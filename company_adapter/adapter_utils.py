"""
DataShield AI — Company Adapter Utilities
==========================================
Shared helpers for text extraction, cleaning, and common preprocessing.
Supports: PDF, DOCX, TXT, XLSX, CSV
"""

import re
import os
import json
import unicodedata
from pathlib import Path
from typing import Optional

# ─── Text Extractors ──────────────────────────────────────────────────────────

def extract_text_from_pdf(filepath: str) -> str:
    """Extract all text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        print(f"  [WARN] PDF extraction failed for {filepath}: {e}")
        return ""


def extract_text_from_docx(filepath: str) -> str:
    """Extract paragraphs + table cells from a DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        print(f"  [WARN] DOCX extraction failed for {filepath}: {e}")
        return ""


def extract_text_from_xlsx(filepath: str) -> str:
    """Convert all sheets of an Excel file to text."""
    try:
        import pandas as pd
        xl = pd.ExcelFile(filepath)
        parts = []
        for sheet in xl.sheet_names:
            df = xl.parse(sheet).astype(str)
            # Include header + values
            parts.append(f"Sheet: {sheet}")
            parts.append(" ".join(df.columns.tolist()))
            for _, row in df.iterrows():
                parts.append(" ".join(row.values.tolist()))
        return "\n".join(parts)
    except Exception as e:
        print(f"  [WARN] XLSX extraction failed for {filepath}: {e}")
        return ""


def extract_text_from_csv(filepath: str) -> str:
    """Convert a CSV file to text."""
    try:
        import pandas as pd
        df = pd.read_csv(filepath, dtype=str, on_bad_lines="skip")
        parts = [" ".join(df.columns.tolist())]
        for _, row in df.iterrows():
            parts.append(" ".join(row.fillna("").values.tolist()))
        return "\n".join(parts)
    except Exception as e:
        print(f"  [WARN] CSV extraction failed for {filepath}: {e}")
        return ""


def extract_text_from_file(filepath: str) -> str:
    """
    Auto-detect file type and extract text.
    Returns raw text or empty string on failure.
    """
    ext = Path(filepath).suffix.lower()
    extractors = {
        ".pdf":  extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".doc":  extract_text_from_docx,
        ".xlsx": extract_text_from_xlsx,
        ".xls":  extract_text_from_xlsx,
        ".csv":  extract_text_from_csv,
        ".txt":  lambda p: open(p, "r", encoding="utf-8", errors="ignore").read(),
        ".md":   lambda p: open(p, "r", encoding="utf-8", errors="ignore").read(),
    }
    extractor = extractors.get(ext)
    if extractor is None:
        print(f"  [WARN] Unsupported file type: {ext}")
        return ""
    return extractor(filepath)


# ─── Text Cleaning ────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Normalize whitespace, remove control characters, keep multilingual chars."""
    # Normalize unicode
    text = unicodedata.normalize("NFC", text)
    # Remove non-printable control chars (keep newlines/tabs)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse repeated whitespace within lines, preserve line breaks
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    lines = [l for l in lines if l]  # drop blank lines
    return "\n".join(lines)


def split_into_sentences(text: str, max_chars: int = 300) -> list[str]:
    """
    Split text into sentence-like chunks suitable for NER.
    Keeps sentences short enough for the spacy model.
    """
    # Split on sentence boundaries
    sentences = re.split(r"(?<=[.!?\n])\s+", text)
    result = []
    for s in sentences:
        s = s.strip()
        if not s:
            continue
        # If sentence is too long, chunk it
        if len(s) > max_chars:
            words = s.split()
            chunk = []
            for word in words:
                chunk.append(word)
                if len(" ".join(chunk)) >= max_chars:
                    result.append(" ".join(chunk))
                    chunk = []
            if chunk:
                result.append(" ".join(chunk))
        else:
            result.append(s)
    return result


# ─── Common Word Filters ──────────────────────────────────────────────────────

# Common English/French/Arabic stopwords to skip in keyword extraction
STOPWORDS = {
    # English
    "the", "and", "for", "that", "this", "with", "from", "have", "will",
    "been", "they", "their", "company", "also", "which", "when", "where",
    "should", "could", "would", "about", "some", "than", "other", "into",
    "more", "these", "those", "each", "such", "after", "before", "over",
    "under", "between", "through", "during", "without", "within", "across",
    "document", "page", "section", "information", "please", "note", "date",
    "name", "email", "phone", "address", "number", "please", "regards",
    # French
    "les", "des", "une", "pour", "dans", "par", "sur", "avec", "que",
    "qui", "est", "son", "ses", "elle", "ils", "elles", "leur", "leurs",
    "cette", "comme", "mais", "tout", "plus", "bien", "aussi", "même",
    # Common document noise
    "confidential", "private", "internal", "draft", "version", "approved",
    "january", "february", "march", "april", "may", "june", "july",
    "august", "september", "october", "november", "december",
    "monday", "tuesday", "wednesday", "thursday", "friday",
}


def is_meaningful_term(term: str, min_len: int = 4) -> bool:
    """
    Return True if a term is worth keeping as a sensitive keyword.
    Filters out common words, pure numbers, single chars, etc.
    """
    t = term.strip().lower()
    if len(t) < min_len:
        return False
    if t in STOPWORDS:
        return False
    # Skip pure numeric strings
    if re.fullmatch(r"[\d\s\-./,]+", t):
        return False
    # Skip URLs and emails (handled by regex engine)
    if re.search(r"https?://|@|www\.", t):
        return False
    return True


# ─── Profile I/O ─────────────────────────────────────────────────────────────

def save_json(data: dict, path: str) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"  Saved → {path}")


def load_json(path: str) -> Optional[dict]:
    if not os.path.exists(path):
        print(f"  [WARN] File not found: {path}")
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ─── Internal Code Pattern Detection ─────────────────────────────────────────

# Patterns that look like internal company codes / identifiers
INTERNAL_CODE_PATTERNS = [
    # PROJ-1234, TASK_ABC, CLIENT-XK2, OP-99
    r"\b[A-Z]{2,8}[-_][A-Z0-9]{1,8}\b",
    # ProjectHelios, OperationSunrise (CamelCase compound words)
    r"\b[A-Z][a-z]{2,}[A-Z][a-z]{2,}(?:[A-Z][a-z]{2,})?\b",
    # V2, v3.1, v12 (version tags attached to project names)
    r"\b[A-Z][A-Za-z0-9]+[-_]?[Vv]\d+(?:\.\d+)?\b",
    # ALL_CAPS_UNDERSCORED (constants, env variables, internal refs)
    r"\b[A-Z]{3,}(?:_[A-Z]{2,}){1,}\b",
]

COMPILED_INTERNAL_PATTERNS = [re.compile(p) for p in INTERNAL_CODE_PATTERNS]


def find_internal_codes(text: str) -> list[str]:
    """Extract tokens that look like internal codes / identifiers."""
    found = set()
    for pattern in COMPILED_INTERNAL_PATTERNS:
        for match in pattern.finditer(text):
            token = match.group(0).strip()
            if len(token) > 3 and not token.lower() in STOPWORDS:
                found.add(token)
    return sorted(found)
